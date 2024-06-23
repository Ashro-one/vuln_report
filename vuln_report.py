import pandas as pd
from docx import Document
import argparse

# 设置命令行参数解析
parser = argparse.ArgumentParser(description='将Excel漏洞数据转换为Word文档')
parser.add_argument('-f', '--file', required=True, help='输入的Excel文件路径')
parser.add_argument('-o', '--output', required=True, help='输出的Word文件路径')
args = parser.parse_args()

# 读取Excel文件
df = pd.read_excel(args.file)

# 创建一个新的Word文档
doc = Document()


# 定义一个函数来添加加粗的标题
def add_bold_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


# 遍历DataFrame的每一行
for index, row in df.iterrows():
    add_bold_paragraph(doc, f"{index + 1}. {row['系统名称']}")

    add_bold_paragraph(doc, "漏洞描述：")
    doc.add_paragraph(f"{row['漏洞描述']}")

    add_bold_paragraph(doc, "漏洞地址：")
    doc.add_paragraph(f"{row['漏洞地址']}")

    add_bold_paragraph(doc, "测试过程")
    doc.add_paragraph(f"{row['测试过程']}")

    # 漏洞等级和等级值放在同一个段落中
    p = doc.add_paragraph()
    run = p.add_run("漏洞等级：")
    run.bold = True
    p.add_run(f"{row['漏洞等级']}")

    add_bold_paragraph(doc, "修复建议：")
    # 按行分割修复建议并添加到文档中
    for suggestion in row['修复建议'].split('\n'):
        doc.add_paragraph(suggestion)
    doc.add_paragraph()  # 在条目之间添加一个空行

# 保存Word文档
doc.save(args.output)
print(f"Word文档已保存到 {args.output}")
